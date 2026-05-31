#!/usr/bin/env python3
"""
RICHGYAN Receipt Generator - Production Web App
Flask backend: DOCX + PDF output, Groq AI autofill
"""

import os
import json
import threading
import requests
from io import BytesIO
from datetime import datetime
from flask import Flask, request, jsonify, send_file, render_template

import zipfile

from docx import Document
from docx.shared import Pt, RGBColor

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# ─── Config ────────────────────────────────────────────────────────────────────
BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, 'RICHGYAN_receipt.docx')
COUNTER_FILE  = os.path.join(BASE_DIR, 'receipt_counter.json')
RECEIPTS_DIR  = os.path.join(BASE_DIR, 'receipts')
GROQ_API_URL  = "https://api.groq.com/openai/v1/chat/completions"

os.makedirs(RECEIPTS_DIR, exist_ok=True)
counter_lock = threading.Lock()

# ─── Counter ───────────────────────────────────────────────────────────────────
def load_counter():
    if os.path.exists(COUNTER_FILE):
        try:
            with open(COUNTER_FILE, 'r') as f:
                return json.load(f).get('counter', 0)
        except Exception:
            return 0
    return 0

def save_counter(val):
    with open(COUNTER_FILE, 'w') as f:
        json.dump({'counter': val}, f)

def get_receipt_id():
    with counter_lock:
        c = load_counter() + 1
        save_counter(c)
        return f"RG-{datetime.now().year}-{str(c).zfill(4)}"

# ─── Shared helpers ────────────────────────────────────────────────────────────
def number_to_words(num):
    ones  = ['','One','Two','Three','Four','Five','Six','Seven','Eight','Nine']
    tens  = ['','','Twenty','Thirty','Forty','Fifty','Sixty','Seventy','Eighty','Ninety']
    teens = ['Ten','Eleven','Twelve','Thirteen','Fourteen','Fifteen',
             'Sixteen','Seventeen','Eighteen','Nineteen']

    if num == 0: return 'Zero'

    def cvt(n):
        r = ''
        if n >= 100: r += ones[n//100]+' Hundred '; n %= 100
        if n >= 20:  r += tens[n//10]+' '; r += (ones[n%10]+' ' if n%10 else '')
        elif n >= 10: r += teens[n-10]+' '
        elif n > 0:  r += ones[n]+' '
        return r.strip()

    parts = []
    h = num % 1000;  num //= 1000
    if h: parts.insert(0, cvt(h))
    th = num % 100;  num //= 100
    if th: parts.insert(0, cvt(th)+' Thousand')
    lk = num % 100;  num //= 100
    if lk: parts.insert(0, cvt(lk)+' Lakh')
    cr = num % 100
    if cr: parts.insert(0, cvt(cr)+' Crore')
    return ' '.join(parts).replace('  ',' ').strip()

def format_amount(amount):
    try: return f"\u20b9 {int(amount):,}"
    except: return str(amount)

def validate_fields(data):
    """Returns (name, amt_int, purpose, month, branch) or raises ValueError with dict."""
    errors = {}
    name    = data.get('name','').strip()
    purpose = data.get('purpose','').strip()
    month   = data.get('month','').strip()
    branch  = data.get('branch','').strip()
    if not name:    errors['name']    = 'Name is required'
    if not purpose: errors['purpose'] = 'Purpose is required'
    if not month:   errors['month']   = 'Month is required'
    if not branch:  errors['branch']  = 'Branch is required'
    try:
        amt = int(str(data.get('amount','')).replace(',','').replace('\u20b9','').strip())
        if amt <= 0: raise ValueError
    except Exception:
        errors['amount'] = 'Valid positive amount is required'
        amt = 0
    if errors:
        raise ValueError(errors)
    return name, amt, purpose, month, branch

# ═══════════════════════════════════════════════════════════════════════════════
# DOCX Generator
# ═══════════════════════════════════════════════════════════════════════════════
def create_receipt_docx(name, amount, purpose, month, branch, receipt_id, today):
    amount_fmt   = format_amount(amount)
    amount_words = number_to_words(amount)

    field_map = {
        'Received with thanks from': f'Received with thanks from {name}',
        'Amount \u2026':             f'Amount {amount_fmt}',
        'In Word \u2026':            f'In Word {amount_words}',
        'For \u2026':                f'For {purpose}',
        'Month \u2026':              f'Month {month}',
        'Branch \u2026':             f'Branch {branch}',
    }

    doc = Document(TEMPLATE_PATH)
    money_receipt_idx = None

    for idx, para in enumerate(doc.paragraphs):
        text = para.text

        if 'MONEY RECEIPT' in text:
            money_receipt_idx = idx
            continue

        if money_receipt_idx is not None and idx == money_receipt_idx + 1 and text.strip() == '':
            if para.runs:
                r = para.runs[0]
                r.text = f'Receipt No: {receipt_id}     Date: {today}'
                r.bold = True
                for extra in para.runs[1:]: extra.text = ''
            else:
                run = para.add_run(f'Receipt No: {receipt_id}     Date: {today}')
                run.bold = True
            money_receipt_idx = None
            continue

        for key, replacement in field_map.items():
            if text.startswith(key) or (key in text and '\u2026' in text):
                if para.runs:
                    first = para.runs[0]
                    saved_bold  = first.bold
                    saved_size  = first.font.size
                    saved_color = None
                    try: saved_color = first.font.color.rgb
                    except: pass
                    for r in para.runs: r.text = ''
                    first.text = replacement
                    first.bold = saved_bold if saved_bold is not None else True
                    if saved_size:  first.font.size = saved_size
                    if saved_color: first.font.color.rgb = saved_color
                else:
                    para.add_run(replacement)
                break

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════════
# PDF Generator  (pure Python via ReportLab — no MS Word, no LibreOffice)
# Works on any cloud host: Render, Railway, Heroku, Docker slim, etc.
# ═══════════════════════════════════════════════════════════════════════════════

# ── Cache logo + signature images from the DOCX template (extracted once) ─────
_TEMPLATE_IMAGES: dict = {}

def _load_template_images():
    """Unzip logo (image1) and signature (image2) from the DOCX template."""
    global _TEMPLATE_IMAGES
    if _TEMPLATE_IMAGES:
        return
    try:
        with zipfile.ZipFile(TEMPLATE_PATH, 'r') as z:
            media = sorted(f for f in z.namelist() if f.startswith('word/media/'))
            for i, path in enumerate(media, start=1):
                _TEMPLATE_IMAGES[f'image{i}'] = BytesIO(z.read(path))
    except Exception:
        pass   # continue without images if template is unavailable

_load_template_images()   # executed once at import time


def format_amount_pdf(amount):
    """PDF-safe amount string — uses 'Rs.' because built-in PDF fonts lack ₹."""
    try:
        return f"Rs. {int(amount):,}"
    except Exception:
        return str(amount)


def create_receipt_pdf(name, amount, purpose, month, branch, receipt_id, today):
    """
    Build a receipt PDF with ReportLab matching the exact RICHGYAN receipt layout:
      • Outer border rectangle around the full page content
      • Logo centered at top, all header text centered (serif font)
      • Thin horizontal rule below header block
      • 'MONEY RECEIPT' centered, bold, underlined
      • Receipt No + Date centered below title
      • 6 data fields, each bold with a light divider line + generous spacing
      • Thank-you block: bold teal header + cursive/script teal body lines
      • Authorized Signature right-aligned (image + label)
      • About section with partial italic text
      • Footer drawn on canvas: website | email | phone
    Amount shows 'Rs.' because built-in PDF fonts lack the ₹ glyph; DOCX keeps ₹.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        HRFlowable, Image as RLImage, Table, TableStyle,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    W, H   = A4          # 595.27 × 841.89 pt
    BORDER = 28          # outer rect distance from page edge
    PAD    = 16          # space between border and text column
    LM     = BORDER + PAD + 6
    RM     = BORDER + PAD + 6

    # ── Register script / cursive font ───────────────────────────────────────
    script_font = 'Times-Italic'   # safe fallback if font file absent
    font_path   = os.path.join(BASE_DIR, 'fonts', 'DancingScript.ttf')
    if os.path.exists(font_path):
        try:
            if 'DancingScript' not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont('DancingScript', font_path))
            script_font = 'DancingScript'
        except Exception:
            pass

    # ── Canvas callback: border + footer ────────────────────────────────────
    def draw_page(canvas, doc):
        canvas.saveState()
        # outer border
        canvas.setLineWidth(0.8)
        canvas.setStrokeColor(colors.black)
        canvas.rect(BORDER, BORDER, W - 2*BORDER, H - 2*BORDER)
        # footer separator line
        fy_line = BORDER + 20
        canvas.setLineWidth(0.4)
        canvas.line(BORDER + 8, fy_line, W - BORDER - 8, fy_line)
        # footer text
        canvas.setFont('Helvetica', 8)
        fy = BORDER + 7
        canvas.setFillColor(colors.HexColor('#1155CC'))
        canvas.drawString(BORDER + 12, fy, 'www.richgyan.com')
        email = 'richgyanindia@gmail.com'
        ew = canvas.stringWidth(email, 'Helvetica', 8)
        canvas.drawString(W / 2 - ew / 2, fy, email)
        canvas.setFillColor(colors.black)
        phone = '9749685050'
        pw = canvas.stringWidth(phone, 'Helvetica', 8)
        canvas.drawString(W - BORDER - 12 - pw, fy, phone)
        canvas.restoreState()

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=LM, rightMargin=RM,
        topMargin=BORDER + PAD + 6,
        bottomMargin=BORDER + PAD + 22,   # leave room for footer
    )

    # ── Style factory ────────────────────────────────────────────────────────
    teal  = colors.HexColor('#0070C0')
    brown = colors.HexColor('#974705')

    def ps(n, **kw):
        base = dict(fontName='Times-Roman', fontSize=11, leading=14, spaceAfter=1)
        base.update(kw)
        return ParagraphStyle(n, **base)

    s_co_name = ps('CoN',  fontName='Times-Bold',   fontSize=14, alignment=TA_CENTER, spaceAfter=1)
    s_co_inst = ps('CoI',  fontName='Times-Bold',   fontSize=13, alignment=TA_CENTER, spaceAfter=1)
    s_co_prog = ps('CoP',  fontName='Times-Bold',   fontSize=10, alignment=TA_CENTER,
                   textColor=brown, spaceAfter=1)
    s_co_reg  = ps('CoR',  fontName='Times-Bold',   fontSize=10, alignment=TA_CENTER, spaceAfter=1)
    s_co_init = ps('CoIn', fontName='Times-Bold',   fontSize=10, alignment=TA_CENTER, spaceAfter=2)
    s_title   = ps('Ttl',  fontName='Times-Bold',   fontSize=16, alignment=TA_CENTER, spaceAfter=0)
    s_rid     = ps('RId',  fontName='Times-Bold',   fontSize=10, alignment=TA_CENTER, spaceAfter=0)
    s_field   = ps('Fld',  fontName='Times-Bold',   fontSize=12, leading=16, spaceAfter=0)
    s_ty_hdr  = ps('TyH',  fontName='Times-Bold',   fontSize=12, textColor=teal, spaceAfter=2)
    s_ty_scr  = ps('TyS',  fontName=script_font,    fontSize=11, textColor=teal,
                   leading=16, spaceAfter=1)
    s_sig_lbl = ps('SLb',  fontName='Times-Bold',   fontSize=11, alignment=TA_CENTER, spaceAfter=0)
    s_abt_ttl = ps('AbT',  fontName='Times-Bold',   fontSize=10, spaceAfter=2)
    s_abt_bdy = ps('AbB',  fontName='Times-Roman',  fontSize=9,  leading=13, spaceAfter=5)

    SP  = lambda h: Spacer(1, h)
    HR  = lambda t=0.75, c=colors.black: HRFlowable(width='100%', thickness=t, color=c)
    HRG = lambda: HRFlowable(width='100%', thickness=0.4, color=colors.HexColor('#AAAAAA'))

    # ── Pull cached images ───────────────────────────────────────────────────
    logo_io = _TEMPLATE_IMAGES.get('image1')
    sig_io  = _TEMPLATE_IMAGES.get('image2')

    story = []

    # ── Logo — centered ───────────────────────────────────────────────────────
    if logo_io:
        logo_io.seek(0)
        logo_img = RLImage(logo_io, width=85, height=42)
        logo_tbl = Table([[logo_img]], colWidths=['100%'])
        logo_tbl.setStyle(TableStyle([('ALIGN', (0,0), (0,0), 'CENTER')]))
        story.append(logo_tbl)
        story.append(SP(4))

    # ── Company header (all centered) ────────────────────────────────────────
    story.append(Paragraph('RICHGYAN INDIA', s_co_name))
    story.append(Paragraph('COMPUTER &amp; VOCATIONAL INSTITUTE', s_co_inst))
    story.append(Paragraph('(An I.T &amp; Vocational Training Awareness Programme)', s_co_prog))
    story.append(Paragraph('Registered under Ministry Of Corporate Affairs. GOVT. OF INDIA', s_co_reg))
    story.append(Paragraph('Initiated By: Richgyan India I.T. &amp; Educational', s_co_init))
    story.append(HR(0.75))
    story.append(SP(22))

    # ── MONEY RECEIPT (underlined) ────────────────────────────────────────────
    story.append(Paragraph('<u>MONEY RECEIPT</u>', s_title))
    story.append(SP(10))

    # ── Receipt No + Date (centered) ─────────────────────────────────────────
    story.append(Paragraph(
        f'Receipt No: {receipt_id}&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;Date: {today}', s_rid,
    ))
    story.append(SP(18))

    # ── Data fields ───────────────────────────────────────────────────────────
    amt_pdf   = format_amount_pdf(amount)
    amt_words = number_to_words(amount)

    for label, value in [
        ('Received with thanks from', name),
        ('Amount', amt_pdf),
        ('In Word', amt_words),
        ('For', purpose),
        ('Month', month),
        ('Branch', branch),
    ]:
        story.append(Paragraph(f'{label}  {value}', s_field))
        story.append(HRG())
        story.append(SP(14))

    story.append(SP(10))

    # ── Thank-you block ───────────────────────────────────────────────────────
    story.append(Paragraph('Thank you for your payment!', s_ty_hdr))
    for msg in [
        'We sincerely appreciate your timely payment and continued trust in our services.',
        'Your support is highly valued and helps us serve you better.',
        'We look forward to continuing our professional relationship.',
        'Thank you once again for your cooperation and confidence in us',
    ]:
        story.append(Paragraph(msg, s_ty_scr))

    story.append(SP(6))

    # ── Signature — right-aligned via Table ──────────────────────────────────
    sig_block = []
    if sig_io:
        sig_io.seek(0)
        sig_img = RLImage(sig_io, width=88, height=35)
        sig_tbl_inner = Table([[sig_img]], colWidths=['100%'])
        sig_tbl_inner.setStyle(TableStyle([('ALIGN', (0,0), (0,0), 'CENTER')]))
        sig_block.append(sig_tbl_inner)
    sig_block.append(Paragraph('<b>Authorized Signature</b>', s_sig_lbl))

    sig_outer = Table([['', sig_block]], colWidths=['55%', '45%'])
    sig_outer.setStyle(TableStyle([
        ('VALIGN',        (0, 0), (-1, -1), 'BOTTOM'),
        ('ALIGN',         (1, 0), (1,  0),  'CENTER'),
        ('LEFTPADDING',   (1, 0), (1,  0),  0),
        ('RIGHTPADDING',  (1, 0), (1,  0),  0),
    ]))
    story.append(sig_outer)
    story.append(SP(12))

    # ── About section ─────────────────────────────────────────────────────────
    story.append(HR(0.75))
    story.append(SP(6))
    story.append(Paragraph('<b>About Richgyan India</b>', s_abt_ttl))
    story.append(Paragraph(
        'Richgyan India has been conducting Training Classes from last 7 years Since 2017. '
        'Richgyan India Computer &amp; Vocational Education.',
        s_abt_bdy,
    ))
    story.append(Paragraph(
        'Richgyan India is Incorporation by Govt. of India. Richgyan Internationally ISO '
        'Certified Organisation, IAF- "International Accreditation Forum" Washington D.C (USA). '
        '<i>Registered under Govt. of India Trade Marks Act, 1999 Section 23 (2), Rule 56 (1).</i>',
        s_abt_bdy,
    ))

    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════════
# Groq AI
# ═══════════════════════════════════════════════════════════════════════════════
def call_groq(api_key, prompt_text):
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    body = {
        "model": "llama3-8b-8192",
        "messages": [
            {"role": "system", "content": (
                "You are a receipt assistant for RICHGYAN INDIA, a computer & vocational institute. "
                "Given a user description, extract receipt fields as JSON. "
                "Respond ONLY with a raw JSON object, no markdown, no explanation. "
                "Fields: name (string), amount (integer in rupees, no symbol), "
                "purpose (string, e.g. 'Basic Computer Course Fee'), "
                "month (string, e.g. 'June 2026'), branch (string, e.g. 'Guwahati'). "
                "If a field cannot be determined, use an empty string or 0."
            )},
            {"role": "user", "content": prompt_text}
        ],
        "temperature": 0.2, "max_tokens": 300,
    }
    resp = requests.post(GROQ_API_URL, headers=headers, json=body, timeout=15)
    resp.raise_for_status()
    content = resp.json()["choices"][0]["message"]["content"].strip()
    if content.startswith("```"):
        content = content.split("```")[1]
        if content.startswith("json"): content = content[4:]
    return json.loads(content)

# ═══════════════════════════════════════════════════════════════════════════════
# Routes
# ═══════════════════════════════════════════════════════════════════════════════
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/groq-autofill', methods=['POST'])
def groq_autofill():
    data    = request.get_json(force=True)
    api_key = data.get('api_key','').strip()
    prompt  = data.get('prompt','').strip()
    if not api_key: return jsonify({'error': 'Groq API key is required'}), 400
    if not prompt:  return jsonify({'error': 'Prompt text is required'}), 400
    try:
        fields = call_groq(api_key, prompt)
        return jsonify({'success': True, 'fields': fields})
    except requests.exceptions.HTTPError as e:
        return jsonify({'error': f'Groq API error: {e.response.status_code} – check your API key'}), 400
    except json.JSONDecodeError:
        return jsonify({'error': 'Groq returned unexpected format'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate', methods=['POST'])
def generate():
    """Generate DOCX receipt."""
    data = request.get_json(force=True)
    try:
        name, amt, purpose, month, branch = validate_fields(data)
    except ValueError as e:
        return jsonify({'error': 'Validation failed', 'fields': e.args[0]}), 400
    try:
        receipt_id = get_receipt_id()
        today      = datetime.now().strftime('%d %B %Y')
        buf        = create_receipt_docx(name, amt, purpose, month, branch, receipt_id, today)
        return send_file(buf, as_attachment=True,
            download_name=f"{receipt_id}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({'error': f'DOCX generation failed: {str(e)}'}), 500

@app.route('/api/generate-pdf', methods=['POST'])
def generate_pdf():
    """Generate PDF receipt."""
    data = request.get_json(force=True)
    try:
        name, amt, purpose, month, branch = validate_fields(data)
    except ValueError as e:
        return jsonify({'error': 'Validation failed', 'fields': e.args[0]}), 400
    try:
        receipt_id = get_receipt_id()
        today      = datetime.now().strftime('%d %B %Y')
        buf        = create_receipt_pdf(name, amt, purpose, month, branch, receipt_id, today)
        return send_file(buf, as_attachment=True,
            download_name=f"{receipt_id}.pdf",
            mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': f'PDF generation failed: {str(e)}'}), 500

@app.route('/api/generate-both', methods=['POST'])
def generate_both():
    """Generate both DOCX and PDF, return JSON with base64 encoded files."""
    import base64
    data = request.get_json(force=True)
    try:
        name, amt, purpose, month, branch = validate_fields(data)
    except ValueError as e:
        return jsonify({'error': 'Validation failed', 'fields': e.args[0]}), 400
    try:
        receipt_id = get_receipt_id()
        today      = datetime.now().strftime('%d %B %Y')
        docx_buf   = create_receipt_docx(name, amt, purpose, month, branch, receipt_id, today)
        pdf_buf    = create_receipt_pdf(name, amt, purpose, month, branch, receipt_id, today)
        return jsonify({
            'receipt_id': receipt_id,
            'date':       today,
            'docx': {
                'filename': f'{receipt_id}.docx',
                'data': base64.b64encode(docx_buf.read()).decode(),
            },
            'pdf': {
                'filename': f'{receipt_id}.pdf',
                'data': base64.b64encode(pdf_buf.read()).decode(),
            }
        })
    except Exception as e:
        return jsonify({'error': f'Generation failed: {str(e)}'}), 500

@app.route('/api/preview', methods=['POST'])
def preview():
    data    = request.get_json(force=True)
    name    = data.get('name','').strip() or '\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026'
    purpose = data.get('purpose','').strip() or '\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026'
    month   = data.get('month','').strip() or '\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026'
    branch  = data.get('branch','').strip() or '\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026'
    today   = datetime.now().strftime('%d %B %Y')
    with counter_lock:
        next_num = load_counter() + 1
    try:
        amt          = int(str(data.get('amount',0)).replace(',','').replace('\u20b9','').strip() or 0)
        amount_fmt   = format_amount(amt) if amt else '\u20b9 \u2026\u2026\u2026\u2026\u2026\u2026'
        amount_words = number_to_words(amt) if amt else '\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026'
    except Exception:
        amount_fmt   = '\u20b9 \u2026\u2026\u2026\u2026\u2026\u2026'
        amount_words = '\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026'
    return jsonify({
        'receipt_id':   f'RG-{datetime.now().year}-{str(next_num).zfill(4)}',
        'date':         today,
        'name':         name,
        'amount_fmt':   amount_fmt,
        'amount_words': amount_words,
        'purpose':      purpose,
        'month':        month,
        'branch':       branch,
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
